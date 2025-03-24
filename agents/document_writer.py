from pathlib import Path
from tempfile import TemporaryDirectory

from docx import Document
import re
from typing import Dict, Any, List, Union
from agentscope.agents import LlamaIndexAgent
from agentscope.message import Msg


class DocumentWriter(LlamaIndexAgent):

    def __init__(
            self,
            name: str,
            model_config_name: str,
            sys_prompt: str,
            knowledge_id_list: List[str],
            recent_n_mem_for_retrieve: int = 3,
    ) -> None:

        sys_prompt = """## Requirements specification generation rules
            You are a professional requirements specification generator, please follow the user input and the rules in the knowledge base:
            1. Generate requirements Specification (SRS) based on existing UML models, requirements decomposition tables and original requirements.
            2. If something is missing in the knowledge base, leave the corresponding part blank.
            3. The RESULT must use the format: ' ' 'RESULT\n[section]: [content]\n... ` ` ` """
        super().__init__(
            name=name,
            sys_prompt=sys_prompt,
            model_config_name=model_config_name,
            knowledge_id_list=knowledge_id_list,
            similarity_top_k=3,
            recent_n_mem_for_retrieve=recent_n_mem_for_retrieve,
        )

    def reply(self, x: Union[Msg, List[Msg]]) -> Msg:

        query = self._extract_query(x)

        related_knowledge = self._retrieve_knowledge(query)

        full_prompt = (
            f"{self.sys_prompt}\n\n"
            f"[Knowledge base content]\n{related_knowledge}\n\n"
            f"[User input]\n{query}\n\n"
            "Generate the requirements specification according to the rules:"
        )

        response_text = self.model(full_prompt).text
        return Msg(self.name, response_text, role="assistant")

    def generate_srs(self, input_data: str) -> Dict[str, Any]:
        response = self.reply(Msg("user", input_data, role="user"))
        return self._parse_response(response.content)

    def _parse_response(self, content: str) -> Dict[str, Any]:
        srs_content = {}
        if match := re.findall(r'```RESULT\n(.*?)\n```', content, re.DOTALL):
            for section in match:
                lines = [line.strip() for line in section.split('\n') if line.strip()]
                for line in lines:
                    if ':' in line:
                        key, value = line.split(':', 1)
                        srs_content[key.strip()] = value.strip()
        return srs_content

    def save_to_doc(self, srs_content: Dict[str, Any], output_path: str) -> None:
        doc = Document()

        self._add_business_level(doc, srs_content)

        self._add_system_level(doc, srs_content)

        self._add_appendixes(doc, srs_content)

        doc.save(output_path)
        print(f"Requirement specifications are generated and saved to {output_path}")

    def _add_heading(self, doc: Document, text: str, level: int):
        doc.add_heading(text, level=level)

    def _add_paragraph(self, doc: Document, text: str):
        doc.add_paragraph(text)

    def _add_table(self, doc: Document, headers: List[str], rows: List[List[str]]):
        table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
        for i, header in enumerate(headers):
            table.cell(0, i).text = header
        for i, row in enumerate(rows):
            for j, cell_value in enumerate(row):
                table.cell(i + 1, j).text = str(cell_value)

    def _add_business_level(self, doc: Document, srs_content: Dict[str, Any]):
        self._add_heading(doc, "Chapter 1: Introduction", level=1)

        # 1.1 Purpose 
        self._add_heading(doc, "1.1 Purpose", level=2)
        purpose = srs_content.get("1.1 Purpose", "")
        self._add_paragraph(doc, purpose if purpose else "Document purpose not provided")

        # 1.2 Glossary 
        self._add_heading(doc, "1.2 Glossary", level=2)
        glossary = srs_content.get("1.2 Glossary", {})
        if glossary:
            headers = ["term", "definition"]
            rows = [[term, definition] for term, definition in glossary.items()]
            self._add_table(doc, headers, rows)
        else:
            self._add_paragraph(doc, "Glossary not provided.")

# 1.3 Stakeholders 
self._add_heading(doc, "1.3 Stakeholder Model", level=2)
stakeholders = srs_content.get("1.3 Stakeholder Model", "")
self._add_paragraph(doc, stakeholders if stakeholders else "No stakeholder model provided.")

# 1.4 Goals Objectives Model
self._add_heading(doc, "1.4 Objectives Model", level=2)
goals = srs_content.get("1.4 Objectives Model", "")
self._add_paragraph(doc, goals if goals else "No objectives model provided.")

# 1.7-1.8 References and Document Structure
self._add_heading(doc, "1.7 References", level=2)
references = srs_content.get("1.7 References", "")
self._add_paragraph(doc, references if references else "No references provided.")

self._add_heading(doc, "1.8 Document Structure", level=2)
structure = srs_content.get("1.8 Document Structure", "")
self._add_paragraph(doc, structure if structure else "No document structure description provided.")

def _add_system_level(self, doc: Document, srs_content: Dict[str, Any]):
    """Add system-level content"""
    self._add_heading(doc, "Chapter 2: System Overview", level=1)

    # 2.1 Context Model
    self._add_heading(doc, "2.1 Context Model", level=2)
    context_model = srs_content.get("2.1 Context Model", "")
    self._add_paragraph(doc, context_model if context_model else "No context model provided.")

    # 2.3 Assumptions
    self._add_heading(doc, "2.3 Assumptions", level=2)
    assumptions = srs_content.get("2.3 Assumptions", "")
    self._add_paragraph(doc, assumptions if assumptions else "No assumptions provided.")

    # 2.4 Constraints
    self._add_heading(doc, "2.4 Constraints", level=2)
    constraints = srs_content.get("2.4 Constraints", "")
    self._add_paragraph(doc, constraints if constraints else "No constraints provided.")

    # Chapters 3~N-1: Subsystem Descriptions
    subsystems = srs_content.get("Subsystem Descriptions", [])
    if subsystems:
        for i, subsystem in enumerate(subsystems, start=1):
            self._add_heading(doc, f"Chapter 3.{i}: Subsystem Description", level=1)
            self._add_heading(doc, f"3.{i}.1 Information Entities", level=2)
            entities = subsystem.get("Information Entities", "")
            self._add_paragraph(doc, entities if entities else "No information entities provided.")

            self._add_heading(doc, f"3.{i}.2 Actors", level=2)
            actors = subsystem.get("Actors", "")
            self._add_paragraph(doc, actors if actors else "No actors provided.")

            self._add_heading(doc, f"3.{i}.3 Functional Requirements", level=2)
            frs = subsystem.get("Functional Requirements", "")
            self._add_paragraph(doc, frs if frs else "No functional requirements provided.")

            self._add_heading(doc, f"3.{i}.4 Use Cases", level=2)
            use_cases = subsystem.get("Use Cases", "")
            self._add_paragraph(doc, use_cases if use_cases else "No use cases provided.")

            self._add_heading(doc, f"3.{i}.5 Subsystem-Level NFRs", level=2)
            nfrs = subsystem.get("Subsystem-Level NFRs", "")
            self._add_paragraph(doc, nfrs if nfrs else "No subsystem-level non-functional requirements provided.")

    # Chapter N: Global Non-functional Requirements
    self._add_heading(doc, "Chapter N: Global Non-functional Requirements", level=1)
    global_nfrs = srs_content.get("Global Non-functional Requirements", "")
    self._add_paragraph(doc, global_nfrs if global_nfrs else "No global non-functional requirements provided.")

    # Chapter N+1: Concerns Integration
    self._add_heading(doc, "Chapter N+1: Concerns Integration", level=1)
    concerns = srs_content.get("Concerns Integration", "")
    self._add_paragraph(doc, concerns if concerns else "No concerns integration content provided.")

def _add_appendixes(self, doc: Document, srs_content: Dict[str, Any]):
    """Add appendix content"""
    self._add_heading(doc, "Appendices", level=1)

    # Document Versions
    self._add_heading(doc, "A. Document Versions", level=2)
    versions = srs_content.get("Document Versions", "")
    self._add_paragraph(doc, versions if versions else "No document version information provided.")

    # Requirements Traceability Matrix
    self._add_heading(doc, "B. Requirements Traceability Matrix", level=2)
    traceability_matrix = srs_content.get("Requirements Traceability Matrix", [])
    if traceability_matrix:
        headers = ["Requirement ID", "Description", "Test Case", "Status"]
        rows = [
            [item["id"], item["description"], item["test_case"], item["status"]]
            for item in traceability_matrix
        ]
        self._add_table(doc, headers, rows)
    else:
        self._add_paragraph(doc, "No requirements traceability matrix provided.")


if __name__ == "__main__":
    ##############################################
    # Example Run (Functional Test)
    ##############################################
    print("\n" + "=" * 50)
    print("Starting example execution of generating Software Requirements Specification...")

    # Initialize generator
    test_writer = DocumentWriter(
        name="Example Generator",
        model_config_name="dummy_model",
        sys_prompt="Testing prompt",
        knowledge_id_list=["sample_knowledge"],
    )

    # Sample input data
    sample_input = """
    [User Requirements]
    1. Need user login functionality
    2. Support file uploads

    [UML Models]
    Class Diagram:
    - User: username, password
    - File: name, size

    [Requirements Breakdown Table]
    FR-001: User login authentication
    NFR-001: File upload response time less than 3 seconds
    """

    # Generate Software Requirements Specification
    print("\nGenerating SRS content...")
    srs_content = test_writer.generate_srs(sample_input)

    # Print generated results
    print("\nGenerated Results Parsing:")
    for section, content in srs_content.items():
        print(f"{section}: {content[:50]}...")  # Show first 50 characters to avoid lengthiness

    # Save test document
    with TemporaryDirectory() as tmp_dir:
        output_path = Path(tmp_dir) / "Example_Document.docx"
        test_writer.save_to_doc(srs_content, str(output_path))
        print(f"\nDocument save verification: {output_path.exists()}")