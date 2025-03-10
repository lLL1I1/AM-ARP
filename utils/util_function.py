import glob
import os

from agentscope.message import Msg
from typing import List, Union
from docx import Document  # 导入 python-docx 库
def _extract_query(x: Union[Msg, List[Msg]]) -> str:
    """提取查询内容"""
    if isinstance(x, Msg):
        return x.content
    elif isinstance(x, list) and len(x) > 0:
        return x[-1].content
    else:
        return ""

def _retrieve_knowledge(query: str, knowledge_list=None, similarity_top_k=None) -> str:
    """执行知识检索"""
    knowledge_str = ""
    for knowledge in knowledge_list:
        nodes = knowledge.retrieve(query, similarity_top_k)
        for node in nodes:
            knowledge_str += f"规则：{node.text}\n来源：{node.node.metadata}\n\n"
    return knowledge_str.strip()

def read_docx(file_path):
    """读取 docx 文件内容并返回字符串"""
    doc = Document(file_path)
    full_text = []
    for paragraph in doc.paragraphs:
        full_text.append(paragraph.text)
    return '\n'.join(full_text)


def get_latest_version_dir() -> Union[str, None]:
    """直接定位到指定目录"""
    # 硬编码绝对路径
    target_dir = r"G:\PycharmProjects\UMLGenerator\agents\output"

    # 验证路径是否存在
    if not os.path.exists(target_dir):
        print(f"[错误] 目录不存在: {target_dir}")
        return None

    # 执行通配符匹配
    pattern = os.path.join(target_dir, "class-*")
    version_dirs = glob.glob(pattern)

    # 过滤并排序
    version_dirs = [d for d in version_dirs if os.path.isdir(d)]
    if version_dirs:
        version_dirs.sort(key=os.path.getctime, reverse=True)
        return version_dirs[0]
    return None


def read_docx(file_path: str) -> str:
    """增强版文档读取，兼容实际docx和文本文件"""
    if file_path.endswith(".docx"):
        # 实际Word文档读取（需要python-docx库）
        from docx import Document
        doc = Document(file_path)
        return "\n".join([para.text for para in doc.paragraphs])
    else:
        # 普通文本文件读取
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()